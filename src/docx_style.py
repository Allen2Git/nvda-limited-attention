'''Shared helpers for generating Chinese academic papers in python-docx.

Implements the standard "国内通用学术论文" formatting:
- Page: A4, top/bottom 2.54cm, left/right 3.17cm
- Title (论文大标题): 小二号 SimHei, centered, bold
- L1 heading (一级标题):  三号 SimHei, left-aligned, bold
- L2 heading (二级标题):  四号 SimHei, left-aligned, bold
- L3 heading (三级标题):  小四 SimHei, left-aligned, bold
- Body:                   小四 SimSun, 1.5 line spacing, first-line indent 2 chars
- Figure/table caption:   五号 SimSun, centered
- References:             五号 SimSun, hanging indent

Critical python-docx trick: font.name only affects Latin; for Chinese we
must also set the rFonts/eastAsia attribute via XML.
'''
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement


# ----- fontsize helpers (Chinese 字号 conventions) ---------------------
SIZE = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24,
    '二号': 22, '小二': 18, '三号': 16, '小三': 15,
    '四号': 14, '小四': 12, '五号': 10.5, '小五': 9, '六号': 7.5,
}


def set_cn_font(run, name_cn: str, name_en: str = 'Times New Roman'):
    '''Set both Latin and East-Asian font of a Run.

    name_cn is the Chinese typeface that appears in Word's font menu
    (e.g. "宋体", "黑体", "仿宋", "楷体").  name_en handles any Latin
    characters in the same run.
    '''
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:cs'), name_en)


# ----- base document factory -------------------------------------------
def enable_line_numbers(section, *, count_by: int = 1, start: int = 1,
                        distance_twips: int = 360,
                        restart: str = 'continuous'):
    '''Turn on continuous line numbers for a section.

    count_by     = show number every N lines  (1 = every line)
    start        = starting number
    distance     = distance from text to numbers, in twips (360 twips = 0.25")
    restart      = 'continuous' | 'newPage' | 'newSection'
    '''
    sectPr = section._sectPr
    # Remove any existing lnNumType to avoid duplicates
    existing = sectPr.find(qn('w:lnNumType'))
    if existing is not None:
        sectPr.remove(existing)
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), str(count_by))
    lnNumType.set(qn('w:start'), str(start))
    lnNumType.set(qn('w:distance'), str(distance_twips))
    lnNumType.set(qn('w:restart'), restart)
    sectPr.append(lnNumType)


def new_doc(line_numbers: bool = True) -> Document:
    doc = Document()
    # Page margins (GB/T 7713 习惯)
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    if line_numbers:
        enable_line_numbers(section, count_by=1, start=1,
                            distance_twips=360, restart='continuous')
    # Default Normal style: 小四 宋体
    normal = doc.styles['Normal']
    normal.font.size = Pt(SIZE['小四'])
    normal.font.name = 'Times New Roman'
    # eastAsia for Normal style
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


# ----- paragraph helpers -----------------------------------------------
def _line_spacing(par, multiple: float = 1.5):
    fmt = par.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple


def _first_line_indent_chars(par, n_chars: int = 2, font_pt: int = 12):
    # Indent = n_chars × font_pt (approximation; Word uses EMUs)
    par.paragraph_format.first_line_indent = Pt(n_chars * font_pt)


def add_title(doc, text: str):
    '''论文大标题：小二黑体，居中，段前段后各 0.5 行。'''
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(12)
    for line in text.split('\n'):
        run = par.add_run(line + '\n')
        run.bold = True
        run.font.size = Pt(SIZE['小二'])
        set_cn_font(run, '黑体', 'Times New Roman')
    # strip trailing newline
    if par.runs and par.runs[-1].text.endswith('\n'):
        par.runs[-1].text = par.runs[-1].text.rstrip('\n')
    return par


def add_subtitle(doc, text: str):
    '''副标题/作者行：小四宋体，居中，斜体。'''
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(12)
    run = par.add_run(text)
    run.italic = True
    run.font.size = Pt(SIZE['小四'])
    set_cn_font(run, '宋体', 'Times New Roman')
    return par


def add_h1(doc, text: str):
    '''一级标题：三号黑体，左对齐，段前 12pt 段后 6pt。'''
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(18)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.line_spacing = 1.5
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(SIZE['三号'])
    set_cn_font(run, '黑体', 'Times New Roman')
    return par


def add_h2(doc, text: str):
    '''二级标题：四号黑体，左对齐，段前 6pt 段后 3pt。'''
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.line_spacing = 1.5
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(SIZE['四号'])
    set_cn_font(run, '黑体', 'Times New Roman')
    return par


def add_h3(doc, text: str):
    '''三级标题：小四黑体，左对齐。'''
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.line_spacing = 1.5
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(SIZE['小四'])
    set_cn_font(run, '黑体', 'Times New Roman')
    return par


def add_body(doc, text: str, *, indent: bool = True, italic: bool = False,
             bold: bool = False, align=None, line_spacing: float = 1.5,
             cn_font: str = '宋体', en_font: str = 'Times New Roman',
             size_key: str = '小四'):
    '''正文段落：小四宋体，1.5 倍行距，首行缩进 2 字符。'''
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    _line_spacing(par, line_spacing)
    if indent:
        par.paragraph_format.first_line_indent = Cm(0.74)  # ≈ 2 个小四字符
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(SIZE[size_key])
    set_cn_font(run, cn_font, en_font)
    return par


def add_caption(doc, text: str):
    '''图/表标题：五号宋体，居中，无缩进，段前 3pt 段后 6pt。'''
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.line_spacing = 1.25
    run = par.add_run(text)
    run.font.size = Pt(SIZE['五号'])
    set_cn_font(run, '宋体', 'Times New Roman')
    return par


def add_reference(doc, text: str):
    '''参考文献条目：五号宋体，悬挂缩进 0.8cm，段后 3pt。'''
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.35
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.8)
    par.paragraph_format.first_line_indent = Cm(-0.8)
    run = par.add_run(text)
    run.font.size = Pt(SIZE['五号'])
    set_cn_font(run, '宋体', 'Times New Roman')
    return par


def style_table(table, *, font_key: str = '五号'):
    '''Apply 宋体 + 五号 to every cell in a table.'''
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.line_spacing = 1.2
                for run in par.runs:
                    run.font.size = Pt(SIZE[font_key])
                    set_cn_font(run, '宋体', 'Times New Roman')
